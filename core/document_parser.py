import os
import re
from pathlib import Path
from datetime import datetime


class DocumentParser:
    """文档解析器基类"""
    
    @staticmethod
    def get_parser(file_path):
        """根据文件类型获取对应的解析器"""
        ext = Path(file_path).suffix.lower()
        
        parsers = {
            '.docx': DocxParser(),
            '.pdf': PdfParser(),
            '.txt': TxtParser(),
            '.xlsx': ExcelParser(),
            '.xls': ExcelParser(),
        }
        
        return parsers.get(ext)
    
    @staticmethod
    def parse(file_path):
        """解析文档，返回文本内容和元数据"""
        parser = DocumentParser.get_parser(file_path)
        if parser:
            return parser.parse(file_path)
        return None, {}
    
    @staticmethod
    def extract_text(file_path):
        """仅提取文本内容"""
        parser = DocumentParser.get_parser(file_path)
        if parser:
            return parser.extract_text(file_path)
        return ""
    
    @staticmethod
    def get_file_info(file_path):
        """获取文件基本信息"""
        path = Path(file_path)
        stat = path.stat()
        
        return {
            'file_name': path.name,
            'file_path': str(path.absolute()),
            'file_type': path.suffix.lower(),
            'file_size': stat.st_size,
            'create_time': datetime.fromtimestamp(stat.st_ctime),
            'modify_time': datetime.fromtimestamp(stat.st_mtime),
        }


class DocxParser:
    """Word文档解析器"""
    
    def parse(self, file_path):
        """解析Word文档"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            
            # 提取文本
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # 提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        paragraphs.append(' | '.join(row_text))
            
            content = '\n'.join(paragraphs)
            
            # 提取元数据
            metadata = {
                'title': doc.core_properties.title or '',
                'author': doc.core_properties.author or '',
                'created': doc.core_properties.created,
                'modified': doc.core_properties.modified,
                'paragraph_count': len(doc.paragraphs),
            }
            
            return content, metadata
            
        except Exception as e:
            print(f"解析Word文档失败: {e}")
            return None, {}
    
    def extract_text(self, file_path):
        content, _ = self.parse(file_path)
        return content or ""


class PdfParser:
    """PDF文档解析器"""
    
    def parse(self, file_path):
        """解析PDF文档"""
        try:
            import fitz  # PyMuPDF
            
            doc = fitz.open(file_path)
            
            paragraphs = []
            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text()
                if text.strip():
                    paragraphs.append(text)
            
            content = '\n'.join(paragraphs)
            
            # 提取元数据
            metadata = {
                'title': doc.metadata.get('title', ''),
                'author': doc.metadata.get('author', ''),
                'page_count': len(doc),
                'creation_date': doc.metadata.get('creationDate', ''),
            }
            
            doc.close()
            
            return content, metadata
            
        except Exception as e:
            print(f"解析PDF文档失败: {e}")
            return None, {}
    
    def extract_text(self, file_path):
        content, _ = self.parse(file_path)
        return content or ""
    
    def extract_images(self, file_path, output_dir=None):
        """提取PDF中的图片"""
        try:
            import fitz
            from PIL import Image
            import io
            
            doc = fitz.open(file_path)
            images = []
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                image_list = page.get_images()
                
                for img_index, img in enumerate(image_list, start=1):
                    xref = img[0]
                    base_image = doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    image_ext = base_image["ext"]
                    
                    if output_dir:
                        output_path = Path(output_dir) / f"page{page_num+1}_img{img_index}.{image_ext}"
                        with open(output_path, 'wb') as f:
                            f.write(image_bytes)
                        images.append(str(output_path))
                    else:
                        images.append(image_bytes)
            
            doc.close()
            return images
            
        except Exception as e:
            print(f"提取PDF图片失败: {e}")
            return []


class TxtParser:
    """纯文本解析器"""
    
    def parse(self, file_path):
        """解析文本文件"""
        try:
            # 尝试检测编码
            import chardet
            
            with open(file_path, 'rb') as f:
                raw_data = f.read()
                result = chardet.detect(raw_data)
                encoding = result['encoding'] or 'utf-8'
            
            content = raw_data.decode(encoding, errors='ignore')
            
            metadata = {
                'encoding': encoding,
                'line_count': len(content.splitlines()),
                'char_count': len(content),
            }
            
            return content, metadata
            
        except Exception as e:
            print(f"解析文本文件失败: {e}")
            return None, {}
    
    def extract_text(self, file_path):
        content, _ = self.parse(file_path)
        return content or ""


class ExcelParser:
    """Excel文档解析器"""
    
    def parse(self, file_path):
        """解析Excel文档"""
        try:
            from openpyxl import load_workbook
            
            wb = load_workbook(file_path, data_only=True)
            
            all_sheets_content = []
            sheet_info = []
            
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                sheet_content = []
                
                for row in sheet.iter_rows(values_only=True):
                    row_text = []
                    for cell in row:
                        if cell is not None:
                            row_text.append(str(cell))
                    if row_text:
                        sheet_content.append(' | '.join(row_text))
                
                if sheet_content:
                    all_sheets_content.append(f"=== {sheet_name} ===")
                    all_sheets_content.extend(sheet_content)
                
                sheet_info.append({
                    'name': sheet_name,
                    'rows': sheet.max_row,
                    'cols': sheet.max_column,
                })
            
            content = '\n'.join(all_sheets_content)
            
            metadata = {
                'sheet_count': len(wb.sheetnames),
                'sheets': sheet_info,
            }
            
            wb.close()
            
            return content, metadata
            
        except Exception as e:
            print(f"解析Excel文档失败: {e}")
            return None, {}
    
    def extract_text(self, file_path):
        content, _ = self.parse(file_path)
        return content or ""


class ArchiveAnalyzer:
    """归档分析器 - 自动识别文档年份和类型"""
    
    YEAR_PATTERNS = [
        r'(\d{4})年',
        r'(\d{4})[-/]',
        r'\[(\d{4})\]',
        r'\((\d{4})\)',
        r'^(\d{4})',
    ]
    
    DOC_TYPE_KEYWORDS = {
        '通知': ['通知', '通告', '通报', '告知'],
        '决定': ['决定', '决议', '决策'],
        '报告': ['报告', '汇报', '总结', '述职'],
        '请示': ['请示', '申请', '呈请'],
        '批复': ['批复', '批示', '回复'],
        '函': ['函', '信函', '来函', '复函'],
        '会议纪要': ['纪要', '会议', '记录', '备忘录'],
        '计划': ['计划', '规划', '方案', '安排'],
        '意见': ['意见', '建议', '指导'],
    }
    
    @classmethod
    def extract_year(cls, text, file_name=None):
        """从文本或文件名中提取年份"""
        sources = []
        if text:
            sources.append(text[:5000])  # 只检查前5000字符
        if file_name:
            sources.append(file_name)
        
        for source in sources:
            for pattern in cls.YEAR_PATTERNS:
                matches = re.findall(pattern, source)
                for match in matches:
                    year = int(match)
                    # 合理的年份范围
                    if 1990 <= year <= 2030:
                        return year
        return None
    
    @classmethod
    def detect_document_type(cls, text, file_name=None):
        """检测文档类型"""
        scores = {}
        
        # 检查文件名
        if file_name:
            file_name_lower = file_name.lower()
            for doc_type, keywords in cls.DOC_TYPE_KEYWORDS.items():
                for keyword in keywords:
                    if keyword in file_name_lower:
                        scores[doc_type] = scores.get(doc_type, 0) + 3
        
        # 检查文本内容（前3000字符）
        if text:
            text_lower = text[:3000].lower()
            for doc_type, keywords in cls.DOC_TYPE_KEYWORDS.items():
                for keyword in keywords:
                    count = text_lower.count(keyword)
                    if count > 0:
                        scores[doc_type] = scores.get(doc_type, 0) + count
        
        if scores:
            return max(scores, key=scores.get)
        return None
    
    @classmethod
    def analyze(cls, file_path, content=None):
        """分析文档，返回年份和类型"""
        file_name = Path(file_path).name
        
        if content is None:
            content = DocumentParser.extract_text(file_path)
        
        year = cls.extract_year(content, file_name)
        doc_type = cls.detect_document_type(content, file_name)
        
        return {
            'year': year,
            'document_type': doc_type,
            'file_name': file_name,
        }


# 便捷函数
def parse_document(file_path):
    """解析文档的便捷函数"""
    return DocumentParser.parse(file_path)


def extract_text(file_path):
    """提取文本的便捷函数"""
    return DocumentParser.extract_text(file_path)


def analyze_for_archive(file_path, content=None):
    """分析文档用于归档的便捷函数"""
    return ArchiveAnalyzer.analyze(file_path, content)
